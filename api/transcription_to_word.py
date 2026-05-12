"""
Transcription-to-Word tool.

Converts raw transcription text into a structured meeting Word document
using a two-step LLM pipeline via the local Ollama server.

Step 1: Transcription → Structured JSON (via LLM)
Step 2: Structured JSON → Word document (via LLM with tool calling)
"""

import json
import os
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_ollama import ChatOllama
from loguru import logger

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

OLLAMA_BASE_URL = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434")
OLLAMA_MODEL = os.environ.get("OLLAMA_MODEL", "mistral")

# ---------------------------------------------------------------------------
# Dutch system prompts
# ---------------------------------------------------------------------------

SYSTEM_PROMPT_STEP1 = """\
Je bent een AI-assistent die vergadertranscripties omzet naar gestructureerde JSON.

Analyseer de onderstaande transcriptietekst en produceer ALLEEN een geldig JSON-object met exact deze structuur:

{
  "titel": "string",
  "aanwezigen": ["string"],
  "datum": "string",
  "starttijd": "uu:mm",
  "eindtijd": "uu:mm",
  "notulist": "string",
  "agenda": ["string"],
  "samenvatting": [
    {
      "hoofdstuk": "string",
      "subtitel": "string",
      "tekst": "string"
    }
  ],
  "actiepunten": [
    {
      "actie": "string",
      "eigenaar": "string",
      "deadline": "string"
    }
  ]
}

Regels:
- Als een waarde onbekend is of niet afgeleid kan worden, gebruik dan "onbekend".
- Voor arrays zonder informatie, gebruik een lege array [].
- Geef ALLEEN geldig JSON terug, zonder markdown-opmaak, backticks of andere tekst.
- De inhoud moet in het Nederlands zijn.
- Vat de besproken onderwerpen samen in logische hoofdstukken.
- Identificeer concrete actiepunten met eigenaar en deadline waar mogelijk.
"""

SYSTEM_PROMPT_SUMMARY_TO_JSON = """\
Je bent een AI-assistent die een tekstvergadering omzet naar gestructureerde JSON.

Je ontvangt een bestaande samenvatting van een vergadering (al opgemaakt in tekst).
Zet deze samenvatting om naar het volgende JSON-formaat:

{
  "titel": "string",
  "aanwezigen": ["string"],
  "datum": "string",
  "starttijd": "uu:mm",
  "eindtijd": "uu:mm",
  "notulist": "string",
  "agenda": ["string"],
  "samenvatting": [
    {
      "hoofdstuk": "string",
      "subtitel": "string",
      "tekst": "string"
    }
  ],
  "actiepunten": [
    {
      "actie": "string",
      "eigenaar": "string",
      "deadline": "string"
    }
  ]
}

Regels:
- Gebruik de inhoud van de samenvatting zo volledig mogelijk — maak niets opnieuw op.
- Als een waarde onbekend is, gebruik dan "onbekend".
- Geef ALLEEN geldig JSON terug, zonder markdown-opmaak of andere tekst.
"""

SYSTEM_PROMPT_STEP2 = """\
Je bent een AI-assistent die een Word-document vult met vergadernotulen.

Je ontvangt een JSON-object met vergadergegevens in deze structuur:
- titel: titel van de vergadering
- aanwezigen: lijst van deelnemers
- datum, starttijd, eindtijd: tijdgegevens
- notulist: naam van de notulist
- agenda: lijst van agendapunten
- samenvatting: lijst van hoofdstukken met subtitel en tekst
- actiepunten: lijst van acties met eigenaar en deadline

Je hebt toegang tot tools om het Word-document te vullen. Gebruik de tools om het document op te bouwen.

Regels:
- Gebruik ALLEEN het exacte bestandspad dat in de instructie wordt gegeven.
- Gebruik GEEN ander pad dan het opgegeven pad.
- Bouw het document logisch op: eerst metadata, dan agenda, dan samenvatting, dan actiepunten.
- Schrijf alles in het Nederlands.

Gebruik de beschikbare tools om het document te vullen. Begin nu.
"""

# ---------------------------------------------------------------------------
# Word document manipulation functions (tools for the LLM)
# ---------------------------------------------------------------------------


def _validate_path(path: str, allowed_path: str) -> bool:
    """Validate that the path exactly matches the allowed template path."""
    return os.path.normpath(path) == os.path.normpath(allowed_path)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a heading to the document."""
    try:
        doc.add_heading(text, level=level)
    except KeyError:
        # Template doesn't have heading style; use bold paragraph
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    """Add a paragraph to the document."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Add a table to the document."""
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass  # Use default table style if Table Grid not available

    # Header row
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            if i < len(row.cells):
                row.cells[i].text = str(cell_text)


def add_bullet_list(doc: Document, items: list[str]) -> None:
    """Add a bulleted list using inline bullet characters — template-style independent."""
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.add_run("\u2022  " + item)


def _clear_document_body(doc: Document) -> None:
    """Remove all body paragraphs and tables, preserving styles and section properties."""
    body = doc.element.body
    sectPr_tag = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr"
    for child in list(body):
        if child.tag != sectPr_tag:
            body.remove(child)


# ---------------------------------------------------------------------------
# LLM interaction
# ---------------------------------------------------------------------------


def _get_ollama_model(temperature: float = 0.1) -> ChatOllama:
    """Create a ChatOllama instance using project conventions."""
    return ChatOllama(
        base_url=OLLAMA_BASE_URL,
        model=OLLAMA_MODEL,
        temperature=temperature,
        format="json",
    )


def _get_ollama_model_no_format(temperature: float = 0.3) -> ChatOllama:
    """Create a ChatOllama instance without forced JSON format."""
    return ChatOllama(
        base_url=OLLAMA_BASE_URL,
        model=OLLAMA_MODEL,
        temperature=temperature,
    )


def _parse_json_response(response_text: str) -> dict[str, Any]:
    """Parse JSON from LLM response, handling common issues."""
    # Remove markdown code fences if present
    cleaned = response_text.strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    cleaned = cleaned.strip()

    try:
        return json.loads(cleaned)
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse JSON response: {e}")
        logger.debug(f"Raw response: {response_text[:500]}")

        # Try to find JSON object in the response
        match = re.search(r"\{[\s\S]*\}", cleaned)
        if match:
            try:
                return json.loads(match.group())
            except json.JSONDecodeError:
                pass

        raise ValueError(
            f"LLM response is not valid JSON. Parse error: {e}"
        ) from e


def _validate_meeting_json(data: dict[str, Any]) -> dict[str, Any]:
    """Validate and normalize the meeting JSON structure."""
    required_fields = [
        "titel", "aanwezigen", "datum", "starttijd", "eindtijd",
        "notulist", "agenda", "samenvatting", "actiepunten"
    ]

    for field in required_fields:
        if field not in data:
            if field in ("aanwezigen", "agenda", "samenvatting", "actiepunten"):
                data[field] = []
            else:
                data[field] = "onbekend"

    # Validate samenvatting items
    if isinstance(data.get("samenvatting"), list):
        for item in data["samenvatting"]:
            if isinstance(item, dict):
                item.setdefault("hoofdstuk", "onbekend")
                item.setdefault("subtitel", "onbekend")
                item.setdefault("tekst", "onbekend")

    # Validate actiepunten items
    if isinstance(data.get("actiepunten"), list):
        for item in data["actiepunten"]:
            if isinstance(item, dict):
                item.setdefault("actie", "onbekend")
                item.setdefault("eigenaar", "onbekend")
                item.setdefault("deadline", "onbekend")

    return data


# ---------------------------------------------------------------------------
# Step 1a (alternative): Existing summary text → Structured JSON
# ---------------------------------------------------------------------------


def step1_summary_to_json(summary_text: str) -> dict[str, Any]:
    """Convert an existing text summary to structured meeting JSON via LLM.

    This is a lighter alternative to step1_transcription_to_json when the
    chatbot already produced a human-readable summary in the conversation.
    The LLM only needs to reformat (not re-analyse) the content.
    """
    logger.info("Step 1 (summary path): Converting existing summary to structured JSON...")

    model = _get_ollama_model(temperature=0.0)

    messages = [
        SystemMessage(content=SYSTEM_PROMPT_SUMMARY_TO_JSON),
        HumanMessage(content=summary_text),
    ]

    response = model.invoke(messages)
    response_text = str(response.content)

    logger.debug(f"Summary-to-JSON response length: {len(response_text)} chars")

    meeting_data = _parse_json_response(response_text)
    meeting_data = _validate_meeting_json(meeting_data)

    logger.info(f"Step 1 (summary path) complete. Title: {meeting_data.get('titel', 'onbekend')}")
    return meeting_data


# ---------------------------------------------------------------------------
# Step 1: Transcription → Structured JSON
# ---------------------------------------------------------------------------


def step1_transcription_to_json(transcription_text: str) -> dict[str, Any]:
    """Convert transcription text to structured meeting JSON via LLM."""
    logger.info("Step 1: Converting transcription to structured JSON...")

    model = _get_ollama_model(temperature=0.1)

    messages = [
        SystemMessage(content=SYSTEM_PROMPT_STEP1),
        HumanMessage(content=transcription_text),
    ]

    response = model.invoke(messages)
    response_text = str(response.content)

    logger.debug(f"LLM response length: {len(response_text)} chars")

    # Parse and validate
    meeting_data = _parse_json_response(response_text)
    meeting_data = _validate_meeting_json(meeting_data)

    logger.info(f"Step 1 complete. Meeting title: {meeting_data.get('titel', 'onbekend')}")
    return meeting_data


# ---------------------------------------------------------------------------
# Step 2: Structured JSON → Word document (LLM-driven)
# ---------------------------------------------------------------------------


def _build_word_document(meeting_data: dict[str, Any], template_path: str, output_path: str) -> str:
    """
    Use the LLM to decide how to populate the Word document from meeting JSON.
    The LLM instructs which content to place where; we execute the tools.
    """
    logger.info("Step 2: Building Word document from structured JSON...")

    # Copy template to output path, then clear body content (preserve styles)
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)
    _clear_document_body(doc)

    # Ask LLM how to structure the document
    model = _get_ollama_model_no_format(temperature=0.3)

    instruction = f"""\
Hier is de vergaderdata als JSON:

{json.dumps(meeting_data, ensure_ascii=False, indent=2)}

Geef instructies voor het opbouwen van het Word-document.
Het document moet bevatten:
1. De titel als hoofdkop
2. Metadata (datum, tijd, aanwezigen, notulist)
3. Agenda
4. Samenvatting per hoofdstuk
5. Actiepunten als tabel

Geef je antwoord als een JSON-array van acties. Elke actie heeft:
- "type": "heading" | "paragraph" | "table" | "bullet_list"
- "content": de inhoud (string voor heading/paragraph, array voor bullet_list)
- "level": (optioneel, voor headings, 1 of 2)
- "bold": (optioneel, voor paragraphs)
- "headers": (voor tables, array van kolomnamen)
- "rows": (voor tables, array van arrays)

Geef ALLEEN de JSON-array terug, geen andere tekst.
Het pad van het document is: {output_path}
Gebruik ALLEEN dit pad.
"""

    messages = [
        SystemMessage(content=SYSTEM_PROMPT_STEP2),
        HumanMessage(content=instruction),
    ]

    response = model.invoke(messages)
    response_text = str(response.content)

    # Try to parse LLM instructions
    try:
        actions = _parse_document_actions(response_text)
        _execute_document_actions(doc, actions)
    except Exception as e:
        logger.warning(f"LLM document instructions failed: {e}. Using fallback builder.")
        # doc body is already cleared; build deterministically
        _fallback_build_document(doc, meeting_data)

    doc.save(output_path)
    logger.info(f"Step 2 complete. Document saved to: {output_path}")
    return output_path


def _parse_document_actions(response_text: str) -> list[dict[str, Any]]:
    """Parse the LLM's document action instructions."""
    cleaned = response_text.strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    cleaned = cleaned.strip()

    try:
        actions = json.loads(cleaned)
        if isinstance(actions, list):
            return actions
    except json.JSONDecodeError:
        # Try to find JSON array
        match = re.search(r"\[[\s\S]*\]", cleaned)
        if match:
            try:
                actions = json.loads(match.group())
                if isinstance(actions, list):
                    return actions
            except json.JSONDecodeError:
                pass

    raise ValueError("Could not parse document actions from LLM response")


def _execute_document_actions(doc: Document, actions: list[dict[str, Any]]) -> None:
    """Execute parsed document actions on the Word document."""
    for action in actions:
        action_type = action.get("type", "")
        content = action.get("content", "")

        if action_type == "heading":
            level = action.get("level", 1)
            add_heading(doc, str(content), level=int(level))

        elif action_type == "paragraph":
            bold = action.get("bold", False)
            add_paragraph(doc, str(content), bold=bool(bold))

        elif action_type == "bullet_list":
            items = content if isinstance(content, list) else [str(content)]
            add_bullet_list(doc, [str(i) for i in items])

        elif action_type == "table":
            headers = action.get("headers", [])
            rows = action.get("rows", [])
            if headers and rows:
                add_table(doc, headers, rows)

        else:
            logger.warning(f"Unknown action type: {action_type}")


def _fallback_build_document(doc: Document, data: dict[str, Any]) -> None:
    """Fallback: build document directly from JSON without LLM instructions."""
    # Title
    add_heading(doc, data.get("titel", "Vergadernotulen"), level=1)

    # Metadata
    add_paragraph(doc, f"Datum: {data.get('datum', 'onbekend')}", bold=True)
    add_paragraph(doc, f"Tijd: {data.get('starttijd', 'onbekend')} - {data.get('eindtijd', 'onbekend')}")
    add_paragraph(doc, f"Notulist: {data.get('notulist', 'onbekend')}")

    # Attendees
    aanwezigen = data.get("aanwezigen", [])
    if aanwezigen:
        add_heading(doc, "Aanwezigen", level=2)
        add_bullet_list(doc, aanwezigen)

    # Agenda
    agenda = data.get("agenda", [])
    if agenda:
        add_heading(doc, "Agenda", level=2)
        add_bullet_list(doc, agenda)

    # Summary
    samenvatting = data.get("samenvatting", [])
    if samenvatting:
        add_heading(doc, "Samenvatting", level=2)
        for item in samenvatting:
            if isinstance(item, dict):
                hoofdstuk = item.get("hoofdstuk", "")
                subtitel = item.get("subtitel", "")
                tekst = item.get("tekst", "")
                if hoofdstuk:
                    add_heading(doc, hoofdstuk, level=3)
                if subtitel and subtitel != hoofdstuk:
                    add_paragraph(doc, subtitel, bold=True)
                if tekst:
                    add_paragraph(doc, tekst)

    # Action items
    actiepunten = data.get("actiepunten", [])
    if actiepunten:
        add_heading(doc, "Actiepunten", level=2)
        headers = ["Actie", "Eigenaar", "Deadline"]
        rows = [
            [
                item.get("actie", "onbekend"),
                item.get("eigenaar", "onbekend"),
                item.get("deadline", "onbekend"),
            ]
            for item in actiepunten
            if isinstance(item, dict)
        ]
        if rows:
            add_table(doc, headers, rows)


# ---------------------------------------------------------------------------
# Main function
# ---------------------------------------------------------------------------


def main(
    transcription_text: str,
    template_word_path: str,
    output_dir: str | None = None,
    existing_summary: str | None = None,
) -> str:
    """
    Convert transcription text (or an existing summary) to a structured Word document.

    Args:
        transcription_text: Raw transcription text to process.
        template_word_path: Path to the Word template file.
        output_dir: Directory for output file. Defaults to same directory as template.
        existing_summary: Optional. Pre-existing text summary from the chat conversation.
            When provided, Step 1 (transcription → JSON) is skipped and this text is
            reformatted into the required JSON structure instead, saving one LLM call.

    Returns:
        Path to the generated .docx file.

    Raises:
        FileNotFoundError: If template_word_path does not exist.
        ValueError: If both transcription_text and existing_summary are empty.
    """
    has_summary = existing_summary and existing_summary.strip()
    has_transcription = transcription_text and transcription_text.strip()

    if not has_summary and not has_transcription:
        raise ValueError("Either transcription_text or existing_summary must be provided")

    template_path = Path(template_word_path).resolve()
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")
    if not template_path.suffix.lower() == ".docx":
        raise ValueError(f"Template must be a .docx file: {template_path}")

    # Determine output path
    if output_dir:
        out_dir = Path(output_dir).resolve()
        out_dir.mkdir(parents=True, exist_ok=True)
    else:
        out_dir = template_path.parent

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"notulen_{timestamp}.docx"
    output_path = str(out_dir / output_filename)

    logger.info("Starting transcription-to-Word pipeline")
    logger.info(f"Template: {template_path}")
    logger.info(f"Output: {output_path}")

    # Step 1: produce structured JSON
    if has_summary:
        logger.info("Using existing chat summary — skipping transcription LLM call.")
        meeting_data = step1_summary_to_json(existing_summary)  # type: ignore[arg-type]
    else:
        meeting_data = step1_transcription_to_json(transcription_text)

    # Step 2: JSON → Word document (LLM-driven)
    result_path = _build_word_document(
        meeting_data,
        str(template_path),
        output_path,
    )

    logger.info(f"Pipeline complete. Output: {result_path}")
    return result_path


# ---------------------------------------------------------------------------
# CLI entry point for testing
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import sys

    # Default test: read from file or stdin
    if len(sys.argv) >= 3:
        # input_file = sys.argv[1]
        template_file = sys.argv[1]
        output_directory = sys.argv[2] if len(sys.argv) >= 4 else None

        # with open(input_file, "r", encoding="utf-8") as f:
        #     text = f.read()

        text = "Dit is een transcriptie van een vergadering, geschreven door Job de Vogel op 5 mei 2026. We hebben vergaderd over het Digilab en Job zorgt voor de samenvatting en actiepunten. De aanwezigen waren Job, Jeroen, en Marieke. We hebben gesproken over de volgende onderwerpen: 1) voortgang van het project, 2) planning van de volgende stappen, en 3) verdeling van taken. De vergadering begon om 10:00 en eindigde om 11:30."

        result = main(text, template_file, output_directory)
        print(f"Generated document: {result}")
    else:
        print("Usage: python transcription_to_word.py <input_text_file> <template.docx> [output_dir]")
        print("")
        print("Or use in Python:")
        print('  from api.transcription_to_word import main')
        print('  result = main(transcription_text, "template.docx")')
        sys.exit(1)
